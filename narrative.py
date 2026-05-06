from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt, RGBColor

from frame_export import FrameSnapshot
from temporal_parameters import TemporalResults


@dataclass
class ReportData:
    title: str
    patient_name: str
    patient_id: str
    exam_date: str
    diagnosis: str
    video_name: str
    markers: dict[str, str]
    paragraphs: list[str]
    notes: str
    frame_snapshots: list[FrameSnapshot]
    frame_montage_bytes: bytes | None
    temporal_results: TemporalResults | None
    temporal_infographic_bytes: bytes | None


def combine_sides(right: str, left: str, both_word: str = "bilateralmente") -> str:
    right = (right or "").strip()
    left = (left or "").strip()
    if right and left:
        if right == left:
            return f"{right} {both_word}."
        return f"{right} à direita e {left} à esquerda."
    if right:
        return f"{right} à direita."
    if left:
        return f"{left} à esquerda."
    return ""


def sentence(prefix: str, body: str) -> str:
    if not body:
        return ""
    prefix = prefix.strip()
    body = body.strip()
    if not body.endswith("."):
        body = f"{body}."
    return f"{prefix} {body}".strip()


def build_report_text(form: dict[str, str]) -> list[str]:
    paragraphs: list[str] = []

    gait_intro = (
        f"Apresenta marcha {form['gait_mode']} {form['support_stability']} e largura do passo "
        f"{form['step_width']}."
    )
    paragraphs.append(gait_intro)

    foot_1 = " ".join(
        item
        for item in [
            sentence(
                "O contato inicial ocorre com",
                combine_sides(form["initial_contact_right"], form["initial_contact_left"]),
            ),
            sentence(
                "Na resposta à carga",
                combine_sides(form["loading_response_right"], form["loading_response_left"]),
            ),
            sentence(
                "O primeiro mecanismo de rolamento está",
                combine_sides(form["first_rocker_right"], form["first_rocker_left"]),
            ),
            sentence(
                "Neste instante há",
                combine_sides(
                    f"{form['retrofoot_right']} do retropé e {form['forefoot_right']} do antepé",
                    f"{form['retrofoot_left']} do retropé e {form['forefoot_left']} do antepé",
                ),
            ),
        ]
        if item
    )
    paragraphs.append(foot_1)

    foot_2 = " ".join(
        item
        for item in [
            sentence(
                "O avanço da perna sobre o pé durante a resposta à carga e apoio simples",
                combine_sides(form["leg_advancement_right"], form["leg_advancement_left"]),
            ),
            sentence(
                "No médio apoio",
                combine_sides(form["midstance_right"], form["midstance_left"]),
            ),
            sentence(
                "O segundo mecanismo de rolamento está",
                combine_sides(form["second_rocker_right"], form["second_rocker_left"]),
            ),
        ]
        if item
    )
    paragraphs.append(foot_2)

    foot_3 = " ".join(
        item
        for item in [
            sentence(
                "A flexão plantar no pré-balanço",
                combine_sides(form["plantar_flexion_right"], form["plantar_flexion_left"]),
            ),
            sentence(
                "O desprendimento do calcanhar do solo ocorre",
                combine_sides(form["heel_off_right"], form["heel_off_left"]),
            ),
            sentence(
                "O terceiro mecanismo de rolamento está",
                combine_sides(form["third_rocker_right"], form["third_rocker_left"]),
            ),
        ]
        if item
    )
    paragraphs.append(foot_3)

    foot_4 = " ".join(
        item
        for item in [
            sentence(
                "A progressão do pé no apoio ocorre em",
                combine_sides(form["foot_progression_right"], form["foot_progression_left"]),
            ),
            sentence(
                "No balanço há",
                combine_sides(form["swing_progression_right"], form["swing_progression_left"]),
            ),
        ]
        if item
    )
    paragraphs.append(foot_4)

    foot_5 = sentence(
        "A liberação do pé",
        combine_sides(form["foot_clearance_right"], form["foot_clearance_left"]),
    )
    paragraphs.append(foot_5)

    knee = " ".join(
        item
        for item in [
            sentence(
                "No contato inicial o joelho apresenta",
                combine_sides(form["knee_initial_right"], form["knee_initial_left"]),
            ),
            sentence(
                "Ocorre",
                combine_sides(form["knee_midstance_right"], form["knee_midstance_left"]),
            ),
            sentence(
                "No pré-balanço há",
                combine_sides(form["knee_pre_swing_right"], form["knee_pre_swing_left"]),
            ),
            sentence(
                "Durante o balanço o joelho apresenta",
                combine_sides(form["knee_swing_right"], form["knee_swing_left"]),
            ),
            sentence(
                "No balanço terminal há",
                combine_sides(form["knee_terminal_right"], form["knee_terminal_left"]),
            ),
        ]
        if item
    )
    paragraphs.append(knee)

    hip = " ".join(
        item
        for item in [
            sentence(
                "No contato inicial o quadril apresenta",
                combine_sides(form["hip_initial_right"], form["hip_initial_left"]),
            ),
            sentence(
                "No apoio terminal",
                combine_sides(form["hip_terminal_right"], form["hip_terminal_left"]),
            ),
            sentence(
                "No balanço a flexão do quadril é",
                combine_sides(form["hip_swing_right"], form["hip_swing_left"]),
            ),
            sentence(
                "Há",
                (
                    f"{form['hip_frontal_stance_right']} do quadril direito e "
                    f"{form['hip_frontal_stance_left']} do quadril esquerdo durante o apoio."
                    if form["hip_frontal_stance_right"] != form["hip_frontal_stance_left"]
                    else f"{form['hip_frontal_stance_right']} dos quadris durante o apoio."
                ),
            ),
            sentence(
                "No balanço ocorre",
                combine_sides(form["hip_frontal_swing_right"], form["hip_frontal_swing_left"]),
            ),
            sentence(
                "A coxa apresenta",
                combine_sides(form["hip_rotation_right"], form["hip_rotation_left"]),
            ),
        ]
        if item
    )
    paragraphs.append(hip)

    pelvis = " ".join(
        item
        for item in [
            form["pelvis_tilt"].strip(),
            form["pelvis_obliquity"].strip(),
            form["pelvis_rotation"].strip(),
        ]
        if item
    )
    paragraphs.append(pelvis)

    trunk = " ".join(
        item
        for item in [form["trunk_position"].strip(), form["trunk_lean"].strip(), form["upper_limbs"].strip()]
        if item
    )
    paragraphs.append(trunk)

    return [paragraph.strip() for paragraph in paragraphs if paragraph.strip()]


def build_docx_bytes(data: ReportData) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(data.title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(26, 43, 60)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Laudo observacional de marcha baseado em análise de vídeo")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(9.5)
    subtitle_run.font.color.rgb = RGBColor(90, 99, 112)

    meta_table = document.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.style = "Table Grid"
    metadata = [
        ("Paciente", data.patient_name or "-"),
        ("Registro", data.patient_id or "-"),
        ("Data do exame", data.exam_date or "-"),
        ("Diagnóstico", data.diagnosis or "-"),
        ("Vídeo", data.video_name or "-"),
        ("Eventos marcados", str(sum(1 for value in data.markers.values() if value.strip()))),
    ]
    for row_index in range(3):
        for col_index in range(2):
            label, value = metadata[row_index * 2 + col_index]
            cell = meta_table.cell(row_index, col_index)
            paragraph = cell.paragraphs[0]
            label_run = paragraph.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(9.5)
            value_run = paragraph.add_run(value)
            value_run.font.size = Pt(9.5)

    document.add_paragraph()
    body_heading = document.add_paragraph()
    body_heading_run = body_heading.add_run("Descrição observacional")
    body_heading_run.bold = True
    body_heading_run.font.size = Pt(12.5)
    body_heading_run.font.color.rgb = RGBColor(26, 43, 60)

    for paragraph in data.paragraphs:
        para = document.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(paragraph)
        run.font.size = Pt(10.5)

    if data.notes.strip():
        notes_heading = document.add_paragraph()
        notes_heading_run = notes_heading.add_run("Observações adicionais")
        notes_heading_run.bold = True
        notes_heading_run.font.size = Pt(12.0)
        notes_heading_run.font.color.rgb = RGBColor(26, 43, 60)
        notes_paragraph = document.add_paragraph()
        notes_paragraph.paragraph_format.line_spacing = 1.15
        notes_paragraph.add_run(data.notes).font.size = Pt(10.5)

    if data.temporal_results and data.temporal_infographic_bytes:
        temporal_heading = document.add_paragraph()
        temporal_heading_run = temporal_heading.add_run("Painel temporal do ciclo de marcha")
        temporal_heading_run.bold = True
        temporal_heading_run.font.size = Pt(12.5)
        temporal_heading_run.font.color.rgb = RGBColor(26, 43, 60)
        temporal_caption = document.add_paragraph()
        temporal_caption_run = temporal_caption.add_run(
            "Representação visual da duração do ciclo, distribuição das fases e posicionamento dos eventos marcados."
        )
        temporal_caption_run.italic = True
        temporal_caption_run.font.size = Pt(9)
        temporal_caption_run.font.color.rgb = RGBColor(90, 99, 112)
        document.add_picture(BytesIO(data.temporal_infographic_bytes), width=Inches(6.7))

    if data.frame_montage_bytes:
        document.add_page_break()
        frame_heading = document.add_paragraph()
        frame_heading_run = frame_heading.add_run("Sequência dos frames do ciclo de marcha")
        frame_heading_run.bold = True
        frame_heading_run.font.size = Pt(12.5)
        frame_heading_run.font.color.rgb = RGBColor(26, 43, 60)
        frame_caption = document.add_paragraph()
        frame_caption_run = frame_caption.add_run("Prancha exportada conforme a prévia exibida no aplicativo.")
        frame_caption_run.italic = True
        frame_caption_run.font.size = Pt(9)
        frame_caption_run.font.color.rgb = RGBColor(90, 99, 112)
        document.add_picture(BytesIO(data.frame_montage_bytes), width=Inches(6.7))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def compact_lines(lines: Iterable[str]) -> str:
    return "\n\n".join(line.strip() for line in lines if line.strip())
